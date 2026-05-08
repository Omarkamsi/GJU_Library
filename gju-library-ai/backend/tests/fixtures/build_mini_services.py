import docx

d = docx.Document()
p = d.add_paragraph()
run = p.add_run("1. About the Library")
run.bold = True
d.add_paragraph(
    "The GJU Library serves students, faculty, and staff with print and digital resources."
)
d.add_paragraph("2. الخدمات")
d.add_paragraph("تقدم المكتبة خدمات الإعارة والبحث في الفهرس.")
d.save("tests/fixtures/mini_services.docx")
