"""
Edit "Electronics proj (1).docx" pages 2-4 to match the actual project:
- Component: BZX84C6V2L (Vz=6.2V, Pz_max=300mW, SOT-23)
- Rs = 1 kΩ, Source = 0–20 V DC
- Vmin = 7.2 V, Vmin(stable) = 11.2 V, Vmax = 20 V (source-limited)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, re

doc = Document("/root/Electronics proj (1).docx")

# ── Helper: replace text in a paragraph preserving runs ──────────────────────
def replace_in_para(para, old, new):
    """Replace old with new across all runs of a paragraph (handles split runs)."""
    full = para.text
    if old not in full:
        return False
    # Simplest safe approach: put everything in first run, clear the rest
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True

# ── 1. FIX CHARACTERISTIC TABLE (Table 0) ────────────────────────────────────
# Should reflect BZX84C6V2L: Vz=6.2V, Izt=5mA, Izm=48mA, Pz=300mW
t0 = doc.tables[0]
fixes_t0 = {
    ("Nominal Zener Voltage", "Vz", "6.3", "V"):  "6.2",
    ("Test Current",          "Izt", "49",  "mA"): "5",
    ("Maximum Zener Current", "Izm", "32",  "mA"): "48",
    ("Maximum Power Dissipation", "P", "200", "mW"): "300",
}
for row in t0.rows:
    cells = [c.text.strip() for c in row.cells]
    for (param, sym, old_val, unit), new_val in fixes_t0.items():
        if cells[0] == param and cells[2] == old_val:
            row.cells[2].paragraphs[0].runs[0].text = new_val

# ── 2. FIX STATE 1 POWER CALCULATION (para 70) ───────────────────────────────
# "19.7 Watt (which means it's safe)" → "19.72 mW (well within 300 mW rating)"
para70 = doc.paragraphs[70]
new_calc = (
    "VRs= Vin-Vz= 10 – 6.2=3.8V\n"
    "Is= 3.8/1000= 3.8mA\n"
    "IL= 6.2/10000= 0.62mA\n"
    "Iz=Is-IL = 3.8mA – 0.62mA = 3.18mA\n"
    "Pz=Iz × Vz = 3.18mA × 6.2V = 19.72 mW  (well within 300 mW rating)"
)
if para70.runs:
    para70.runs[0].text = new_calc
    for r in para70.runs[1:]:
        r.text = ""

# ── 3. UPDATE STATE 2 DESCRIPTION ────────────────────────────────────────────
# State 2 currently says RL=500Ω; update to match Rs=1kΩ, 6.2V Zener context
# With Rs=1kΩ and Vin=10V, minimum RL for regulation = Vz/Is = 6.2/3.8e-3 ≈ 1.63kΩ
# So RL=1kΩ is below threshold — update explanation
para74 = doc.paragraphs[74]
replace_in_para(
    para74,
    "500Ω",
    "1Ωk"
)
# Full replacement for accuracy
new_state2 = (
    "In this state, the load resistance is reduced to 1 kΩ, "
    "which falls below the minimum threshold of approximately 1.63 kΩ required to maintain "
    "regulation. With Rs = 1 kΩ and Vin = 10 V, "
    "the maximum available source current Is = 3.8 mA. "
    "A load of 1 kΩ would require IL = 6.2 mA to stay regulated, "
    "which exceeds Is, so the Zener turns off and Vout drops to ≈50 % of Vin ≈ 5 V."
)
if para74.runs:
    para74.runs[0].text = new_state2
    for r in para74.runs[1:]:
        r.text = ""

para75 = doc.paragraphs[75]
new_state2b = (
    "An output of ≈5 V on a Zener rated at 6.2 V means the diode never enters "
    "the breakdown region, defeating the purpose of voltage regulation."
)
if para75.runs:
    para75.runs[0].text = new_state2b
    for r in para75.runs[1:]:
        r.text = ""

# ── 4. UPDATE STATE 3 DESCRIPTION ────────────────────────────────────────────
# Update power reference to match 300 mW rating
para83 = doc.paragraphs[83]
old83 = para83.text
new83 = old83.replace(
    "exceeding its power rating",
    "exceeding its 300 mW power rating (BZX84C6V2L, SOT-23)"
).replace(
    "exceed its power rating",
    "exceed its 300 mW power rating (BZX84C6V2L, SOT-23)"
)
if para83.runs and old83 != new83:
    para83.runs[0].text = new83
    for r in para83.runs[1:]:
        r.text = ""

# ── 5. ADD Vmin/Vmax PARAGRAPH after "Characteristic table:" heading ─────────
# Find para 87 ("Characteristic table:")
# Insert Vmin/Vmax info before the table — we append after para 87
# We'll add it to the Results section (para 91+) instead, as text update

# ── 6. FIX RESULTS SECTION TEXT (para 92, 93) ────────────────────────────────
# Currently says "5.1 V" and "100 Ω" → update to "6.2 V" and "1 kΩ"
para92 = doc.paragraphs[92]
new92 = (
    "The Zener voltage regulator was tested with different load resistances using a "
    "variable load resistor. The input voltage was set to 10 V, the series resistor "
    "Rs = 1 kΩ, and the BZX84C6V2L Zener diode (Vz = 6.2 V, "
    "Pz_max = 300 mW) was used. A DC sweep from 0 V to 20 V was also "
    "performed in LTSpice to determine the input voltage range for breakdown-region operation: "
    "Vmin = 7.2 V and Vmax = 20 V (source-limited)."
)
if para92.runs:
    para92.runs[0].text = new92
    for r in para92.runs[1:]:
        r.text = ""

para93 = doc.paragraphs[93]
new93 = (
    "When in regulation, the voltage across the Zener diode was constant at 6.2 V for all "
    "valid load conditions. The current through the series resistor Is = (Vin−Vz)/Rs "
    "= (10−6.2)/1000 = 3.8 mA. "
    "Zener current Iz varied with load but remained positive, confirming regulation. "
    "Maximum Zener power dissipation at Vmax = 20 V: "
    "Iz = 13.8 mA, Pz = 85.6 mW (28.5 % of 300 mW rating)."
)
if para93.runs:
    para93.runs[0].text = new93
    for r in para93.runs[1:]:
        r.text = ""

# ── 7. FIX RESULTS TABLE (Table 1) ───────────────────────────────────────────
# Replace with values for BZX84C6V2L, Rs=1kΩ, Vin=10V
# New rows: RL=10kΩ (regulated), RL=3kΩ (regulated), RL=1kΩ (unregulated)
t1 = doc.tables[1]

# Header row — already correct ("Load Resistance", "Zener Voltage", "Load Current", "Zener Current")
# Update header to clarify units
header_new = ["Load Resistance (Ω)", "Output Voltage (V)", "Load Current (mA)", "Zener Current (mA)"]
for j, cell in enumerate(t1.rows[0].cells):
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = header_new[j]

# Data rows (rows 1, 2, 3)
data_new = [
    ["10000 (10 kΩ)", "6.2", "0.62", "3.18"],
    ["3000 (3 kΩ)",   "6.2", "2.07", "1.73"],
    ["1000 (1 kΩ)",   "≈ 5.0 (unregulated)", "5.00", "0 (Zener off)"],
]
for i, row_data in enumerate(data_new, start=1):
    row = t1.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = val
        else:
            cell.paragraphs[0].add_run(val)

# ── 8. FIX "For maximum / minimum load resistance" paras (95, 97) ─────────────
para95 = doc.paragraphs[95]
new95 = (
    "For maximum load resistance (RL = 10 kΩ, Vin = 10 V, Rs = 1 kΩ):"
    "\n    Is = (10 – 6.2) / 1000 = 3.8 mA"
    "\n    IL = 6.2 / 10000 = 0.62 mA"
    "\n    Iz = 3.8 – 0.62 = 3.18 mA   → Zener ON, Vout = 6.2 V"
)
if para95.runs:
    para95.runs[0].text = new95
    for r in para95.runs[1:]:
        r.text = ""

para97 = doc.paragraphs[97]
new97 = (
    "For minimum load resistance (RL = 1 kΩ — below regulation threshold):"
    "\n    IL needed = 6.2 / 1000 = 6.2 mA  >  Is = 3.8 mA  → Zener OFF"
    "\n    Vout = Vin × RL / (Rs + RL) = 10 × 1000 / 2000 ≈ 5.0 V"
    "\n    RL_min for regulation = Vz / Is = 6.2 / 0.0038 ≈ 1.63 kΩ"
)
if para97.runs:
    para97.runs[0].text = new97
    for r in para97.runs[1:]:
        r.text = ""

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("/root/Electronics proj (1).docx")
print("Done — saved updated document.")
